from docx2python import docx2python
import docx
import re

from app.main.mse22.document.page_object import PageObjectImage, PageObjectTable, PageObjectHeader
from app.main.mse22.document.page import Page


class PageCreator:
    @staticmethod
    def makePageObject(docx_docx2python, docx_docx, i, j, paragraph_index):
        image_pattern = r'\-{4}media/image\d+\.\D+\-{4}'
        if re.search(image_pattern, docx_docx2python.body[i][0][0][j - 1]):
            return PageObjectImage('image', docx_docx.paragraphs[paragraph_index])
        else:
            return PageObjectHeader('paragraph', docx_docx.paragraphs[paragraph_index])

    @staticmethod
    def make_content(doc, body_index, paragraph_index):
        new_chapters = []
        index = 0
        if len(doc.body[body_index][0][0]) == paragraph_index + 1:
            body_index += 1
            paragraph_index = 0
        for elem in doc.body[body_index]:
            if elem[0][0] == 'ЗАКЛЮЧЕНИЕ':
                break
            elif elem[0][0] == 'ВВЕДЕНИЕ' or elem[0][0] == 'ОПРЕДЕЛЕНИЯ, ОБОЗНАЧЕНИЯ И СОКРАЩЕНИЯ' or elem[0][0] == '':
                continue
            else:
                new_chapters.append(elem[0][0])
        else:
            paragraph_index += 1
            if 'ЗАКЛЮЧЕНИЕ' not in doc.body[body_index][0][0]:
                return False, None, None
            for i in range(paragraph_index, len(doc.body[body_index][0][0])):
                if doc.body[body_index][0][0][i] == 'ВВЕДЕНИЕ' \
                        or doc.body[body_index][0][0][i] == 'ОПРЕДЕЛЕНИЯ, ОБОЗНАЧЕНИЯ И СОКРАЩЕНИЯ' \
                        or doc.body[body_index][0][0][i] == '':
                    continue
                elif doc.body[body_index][0][0][i] == 'ЗАКЛЮЧЕНИЕ':
                    index = i + 1
                    break
                else:
                    new_chapters.append(doc.body[body_index][0][0][i])
        return True, new_chapters, index

    def createPageObjects(self, path, file_type):
        docx_docx2python = docx2python(path)
        indices, errors = self.makeIndices(docx_docx2python, file_type)
        docx_docx = docx.Document(path)

        table_index = 0
        paragraph_index = 0
        page_objects = []
        current_page_object = []

        for i in range(0, len(indices)):
            chapter = docx_docx.paragraphs[paragraph_index].text
            for j in range(indices[i][0][0], indices[i][1][0] + 1):
                if i == 0:
                    chapter = 'Титульный лист'

                if len(docx_docx2python.body[j]) > 1:
                    current_page_object.append(PageObjectTable('table', docx_docx.tables[table_index]))
                    table_index += 1

                else:
                    if indices[i][0][0] == indices[i][1][0]:
                        for _ in range(indices[i][0][1], indices[i][1][1] + 1):
                            current_page_object.append(self.makePageObject(docx_docx2python, docx_docx, j, _,
                                                                           paragraph_index))
                            paragraph_index += 1

                    elif j == indices[i][0][0] and j != indices[i][1][0]:
                        for _ in range(indices[i][0][1], len(docx_docx2python.body[j][0][0])):
                            current_page_object.append(self.makePageObject(docx_docx2python, docx_docx, j, _,
                                                                           paragraph_index))
                            paragraph_index += 1

                    elif j != indices[i][0][0] and j != indices[i][1][0]:
                        for _ in range(0, len(docx_docx2python.body[j][0][0])):
                            current_page_object.append(self.makePageObject(docx_docx2python, docx_docx, j, _,
                                                                           paragraph_index))
                            paragraph_index += 1

                    elif j == indices[i][1][0]:
                        for _ in range(0, indices[i][1][1] + 1):
                            current_page_object.append(self.makePageObject(docx_docx2python, docx_docx, j, _,
                                                                           paragraph_index))
                            paragraph_index += 1

            page_objects.append(Page(chapter, current_page_object))
            current_page_object = []

        return page_objects, errors

    def makeIndices(self, doc_result, file_type):
        if file_type == 'LR':
            chapters = ['Цель работы', 'Основные теоретические положения',
                        'Выполнение работы', 'Тестирование', 'Выводы']
        else:
            chapters = ['ЗАДАНИЕ НА ВЫПУСКНУЮ КВАЛИФИКАЦИОННУЮ РАБОТУ',
                        'КАЛЕНДАРНЫЙ ПЛАН ВЫПОЛНЕНИЯ ВЫПУСКНОЙ КВАЛИФИКАЦИОННОЙ РАБОТЫ', 'РЕФЕРАТ', 'ABSTRACT',
                        'СОДЕРЖАНИЕ', 'ОПРЕДЕЛЕНИЯ, ОБОЗНАЧЕНИЯ И СОКРАЩЕНИЯ', 'ВВЕДЕНИЕ', 'ЗАКЛЮЧЕНИЕ',
                        'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ']
        i_start = [0, 0]
        docx_chapters = []
        errors = []
        cur_index = 0
        cur_chapter = 0
        application_pattern = r'ПРИЛОЖЕНИЕ [ЙЦУКЕНГШЩЗХЪФЫВАПРОЛДЖЭЯЧСМИТЬБЮ]{1}'
        find_index = 0
        while cur_chapter != len(chapters):
            if cur_index == len(doc_result.body):
                if chapters[cur_chapter] != 'ОПРЕДЕЛЕНИЯ, ОБОЗНАЧЕНИЯ И СОКРАЩЕНИЯ':
                    errors.append(chapters[cur_chapter])
                cur_chapter += 1

                if not docx_chapters:
                    cur_index = 0
                else:
                    cur_index = docx_chapters[-1][0][0]

                if cur_chapter == len(chapters):
                    break

            if len(doc_result.body[cur_index]) > 1:  # table
                cur_index += 1

            elif chapters[cur_chapter] in doc_result.body[cur_index][0][0][i_start[1]:]:
                if chapters[cur_chapter] == 'СОДЕРЖАНИЕ':
                    is_correct, new_chapters, index = self.make_content(doc_result, cur_index,
                                                                        doc_result.body[cur_index][0][0]
                                                                        .index(chapters[cur_chapter]))
                    if is_correct:
                        chapters = chapters[:cur_chapter + 3] + new_chapters + chapters[cur_chapter + 3:]

                        if not docx_chapters:
                            docx_chapters.append([i_start, [cur_index, doc_result.body[cur_index][0][0].
                                                 index(chapters[cur_chapter]) - 1]])
                        else:
                            docx_chapters.append([[docx_chapters[-1][1][0], docx_chapters[-1][1][1] + 1],
                                                  [cur_index,
                                                   doc_result.body[cur_index][0][0].index(chapters[cur_chapter]) - 1]])

                        find_index = index
                        i_start = [docx_chapters[-1][1][0], docx_chapters[-1][1][1] + 1]
                        cur_chapter += 1
                        continue

                tmp_i_end = [0, 0]
                tmp_i_start = i_start.copy()
                tmp_i_end[0] = cur_index
                tmp_i_end[1] = doc_result.body[cur_index][0][0].index(chapters[cur_chapter], find_index) - 1
                docx_chapters.append([tmp_i_start, tmp_i_end])
                i_start[0], i_start[1] = tmp_i_end[0], tmp_i_end[1] + 1
                cur_chapter += 1

            else:
                cur_index += 1
        if not docx_chapters:
            docx_chapters.append([[0, 0], [len(doc_result.body) - 1, len(doc_result.body[-1]) - 1]])
        for i in range(cur_index, len(doc_result.body)):
            if len(doc_result.body[i][0][0]) > 1:
                for j in range(docx_chapters[-1][1][1], len(doc_result.body[i][0][0])):
                    if re.search(application_pattern, doc_result.body[i][0][0][j]):
                        docx_chapters.append([[docx_chapters[-1][1][0], docx_chapters[-1][1][1] + 1],
                                              [i, j - 1]])
            cur_index += 1
        docx_chapters.append([[docx_chapters[-1][1][0], docx_chapters[-1][1][1] + 1],
                              [len(doc_result.body) - 1, len(doc_result.body[-1][0][0]) - 1]])
        return docx_chapters, errors
