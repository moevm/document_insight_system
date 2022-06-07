import pytest
import docx
from docx2python import docx2python

from app.main.mse22.document.page_creator import PageCreator


@pytest.mark.parametrize("entities, body, paragraph_in_body, paragraph, expected_type, expected_content", [
    (
        [(1, 0, 0)], 0, 0, 0, "paragraph", 0
    ),
    (
        [(2, 0, 0)], 0, 1, 1, "paragraph", 1
    ),
    (
        [(3, 0, 0)], 0, 1, 1, "paragraph", 1
    ),
    (
        [(1, 1, 0), (3, 0, 0)], 2, 1, 2, "paragraph", 1
    ),
    (
        [(0, 1, 0), (2, 0, 0), (1, 0, 0)], 1, 1, 1, "paragraph", 1
    ),
    (
        [(0, 0, 1)], 0, 0, 0, "image", None
    ),
    (
        [(1, 0, 1), (1, 0, 0)], 0, 1, 2, "paragraph", 0
    ),
    (
        [(1, 0, 1)], 0, 0, 1, "image", None
    ),
    (
        [(1, 1, 1)], 2, 0, 1, "image", None
    ),
    (
        [(0, 1, 1)], 1, 0, 0, "image", None
    )
])
def test_make_page_object(tmp_path, entities, body, paragraph_in_body, paragraph, expected_type, expected_content):
    filename = tmp_path.as_posix() + "/test_document.docx"
    docx_document = docx.Document()
    for pair in entities:
        paragraphs, tables, images = pair
        for i in range(paragraphs):
            docx_document.add_paragraph(str(i))
        for i in range(tables):
            docx_document.add_table(1, 1)
        for i in range(images):
            docx_document.add_picture("../resources/cat.jpg", 432, 540)
    docx_document.save(filename)
    docx2python_document = docx2python(filename)
    page_object = PageCreator.makePageObject(docx2python_document, docx_document,
                                                 body, paragraph_in_body, paragraph)
    assert page_object.type == expected_type
    if page_object.type == "paragraph":
        assert page_object.data.text == str(expected_content)


# Отслеживаемые в PageCreator.make_content() разделы:
checked_sections = ["ВВЕДЕНИЕ", "ОПРЕДЕЛЕНИЯ, ОБОЗНАЧЕНИЯ И СОКРАЩЕНИЯ", "ЗАКЛЮЧЕНИЕ"]


@pytest.mark.parametrize("filename, body, paragraph, expected_result", [
    # Список без раздела "Заключение"
    (
        "../test_files/content/list-no-conclusion.docx", 0, 2, (False, None, None)
    ),
    # Список со всеми требуемыми отслеживаемыми разделами (отслеживаемые разделы в содержании капсом)
    (
        "../test_files/content/list-with-mandatory-caps.docx", 0, 2,
        (True, ["1. Изучение литературы", "1.1. Текст 1", "1.2. Текст 2"], 8)
    ),
    # Список без некоторых требуемых отслеживаемых разделов (присутствующие отслеживаемые разделы в содержании капсом)
    # (нет переноса строки после содержания)
    (
        "../test_files/content/list-without-mandatory.docx", 0, 2,
        (True, ["1. Изучение литературы", "1.1. Текст 1", "1.2. Текст 2"], 7)
    ),
    # Список со всеми требуемыми отслеживаемыми разделами (отслеживаемые разделы в содержании обычным шрифтом)
    (
        "../test_files/content/list-with-mandatory-normal.docx", 0, 2,
        (True, ["1. Изучение литературы", "1.1. Текст 1", "1.2. Текст 2"], 8)
    ),
    # Таблица без раздела "Заключение"
    (
        "../test_files/content/table-no-conclusion.docx", 0, 0,
        (False, None, None)
    ),
    # Таблица со всеми требуемыми отслеживаемыми разделами (отслеживаемые разделы в содержании капсом)
    (
        "../test_files/content/table-with-mandatory-caps.docx", 0, 0,
        (True, ["1. Текущее положение вещей"], 0)
    ),
    # Таблица без некоторых требуемых отслеживаемых разделов (присутствующие отслеживаемые разделы в содержании капсом)
    (
        "../test_files/content/table-without-mandatory.docx", 0, 0,
        (True, ["Текст", "1. Текущее положение вещей"], 0)
    ),
    # Таблица со всеми требуемыми отслеживаемыми разделами (отслеживаемые разделы в содержании обычным шрифтом)
    (
        "../test_files/content/table-with-mandatory-normal.docx", 0, 0,
        (True, ["1. Текущее положение вещей"], 0)
    ),
    # Автосгенерированное содержание
    (
        "../test_files/content/toc.docx", 0, 0,
        (True, ["1. Первый раздел", "1.1. Первый подраздел первого раздела",
                "1.2. Второй подраздел первого раздела", "2. ВТОРОЙ раздел", "2.1. Первый подраздел второго раздела",
                "2.2. Второй подраздел второго раздела"], 0)
    )
])
def test_make_content(filename, body, paragraph, expected_result):
    assert PageCreator.make_content(docx2python(filename), body, paragraph) == expected_result


class TestMakeIndices:
    chapters_lab = ['Цель работы', 'Основные теоретические положения',
                        'Выполнение работы', 'Тестирование', 'Выводы']
    chapters_fwq = ['ЗАДАНИЕ НА ВЫПУСКНУЮ КВАЛИФИКАЦИОННУЮ РАБОТУ',
                        'КАЛЕНДАРНЫЙ ПЛАН ВЫПОЛНЕНИЯ ВЫПУСКНОЙ КВАЛИФИКАЦИОННОЙ РАБОТЫ', 'РЕФЕРАТ', 'ABSTRACT',
                        'СОДЕРЖАНИЕ', 'ОПРЕДЕЛЕНИЯ, ОБОЗНАЧЕНИЯ И СОКРАЩЕНИЯ', 'ВВЕДЕНИЕ', 'ЗАКЛЮЧЕНИЕ',
                        'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ']

    def test_empty_file(self, tmp_path):
        docx_document = docx.Document()
        filename = tmp_path.as_posix() + "/empty.docx"
        docx_document.save(filename)
        docx2python_document = docx2python(filename)
        errors = PageCreator().makeIndices(docx2python_document, "LR")[1]
        assert errors == self.chapters_lab
        errors = PageCreator().makeIndices(docx2python_document, "FWQ")[1]
        assert errors == self.chapters_fwq

    @pytest.mark.parametrize("paragraphs, file_type, expected_chapters, expected_errors", [
        (
            ["", "Цель работы", "Основные теоретические положения", "Выполнение работы", "Тестирование", "Выводы"],
            "LR",
            [
                [[0, 0], [0, 0]],
                [[0, 1], [0, 1]],
                [[0, 2], [0, 2]],
                [[0, 3], [0, 3]],
                [[0, 4], [0, 4]],
                [[0, 5], [0, 5]]
            ],
            []
        ),
        (
                ["Цель работы", "Основные теоретические положения", "Выполнение работы", "Тестирование", "Выводы"],
                "LR",
                [
                    [[0, 0], [0, -1]],
                    [[0, 0], [0, 0]],
                    [[0, 1], [0, 1]],
                    [[0, 2], [0, 2]],
                    [[0, 3], [0, 3]],
                    [[0, 4], [0, 4]]
                ],
                []
        ),
        (
            ["", "Цель работы", "Основные теоретические положения", "Выполнение работы",
             "Тестирование", "Выводы", "ПРИЛОЖЕНИЕ А"],
            "LR",
            [
                [[0, 0], [0, 0]],
                [[0, 1], [0, 1]],
                [[0, 2], [0, 2]],
                [[0, 3], [0, 3]],
                [[0, 4], [0, 4]],
                [[0, 5], [0, 5]],
                [[0, 6], [0, 6]]
            ],
            []
        ),
        (
            ["", "Цель работы", "Основные теоретические положения", "Выполнение работы",
             "Тестирование", "Выводы", "Приложение А"],
            "LR",
            [
                [[0, 0], [0, 0]],
                [[0, 1], [0, 1]],
                [[0, 2], [0, 2]],
                [[0, 3], [0, 3]],
                [[0, 4], [0, 4]],
                [[0, 5], [0, 5]],
                [[0, 6], [0, 6]]
            ],
            []
        ),
        (
                ["", "", "Цель работы", "Основные теоретические положения", "Выполнение работы",
                 "Тестирование", "Выводы", "ПРИЛОЖЕНИЕ А", "ПРИЛОЖЕНИЕ Б", "", "ПРИЛОЖЕНИЕ В"],
                "LR",
                [
                    [[0, 0], [0, 1]],
                    [[0, 2], [0, 2]],
                    [[0, 3], [0, 3]],
                    [[0, 4], [0, 4]],
                    [[0, 5], [0, 5]],
                    [[0, 6], [0, 6]],
                    [[0, 7], [0, 7]],
                    [[0, 8], [0, 9]],
                    [[0, 10], [0, 10]]
                ],
                []
        ),
        (
            ["Text", "Цель работы", "Text", "Основные теоретические положения", "Text", "Выполнение работы",
             "Text", "Тестирование", "Text", "Выводы", "Text"],
            "LR",
            [
                [[0, 0], [0, 0]],
                [[0, 1], [0, 2]],
                [[0, 3], [0, 4]],
                [[0, 5], [0, 6]],
                [[0, 7], [0, 8]],
                [[0, 9], [0, 10]]
            ],
            []
        ),
        (
            ["", "Основные теоретические положения", "Выполнение работы", "Тестирование", ""],
            "LR",
            [
                [[0, 0], [0, 0]],
                [[0, 1], [0, 1]],
                [[0, 2], [0, 2]],
                [[0, 3], [0, 4]],
            ],
            ["Цель работы", "Выводы"]
        ),
        (
            ["", "[table]", "Цель работы", "Основные теоретические положения", "Выполнение работы",
             "Тестирование", "Выводы"],
            "LR",
            [
                [[0, 0], [1, 0]],
                [[2, 0], [2, 0]],
                [[2, 1], [2, 1]],
                [[2, 2], [2, 2]],
                [[2, 3], [2, 3]],
                [[2, 4], [2, 4]]
            ],
            []
        ),
        (
            ["", "[table]", "", "Цель работы", "Основные теоретические положения", "Выполнение работы",
                 "Тестирование", "Выводы"],
            "LR",
            [
                [[0, 0], [2, 0]],
                [[2, 1], [2, 1]],
                [[2, 2], [2, 2]],
                [[2, 3], [2, 3]],
                [[2, 4], [2, 4]],
                [[2, 5], [2, 5]]
            ],
            []
        ),
        (
            ["", "[table]", "", "Цель работы", "Text", "", "Основные теоретические положения", "", "[table]",
             "", "Выполнение работы", "Тестирование", "Выводы", ""],
            "LR",
            [
                [[0, 0], [2, 0]],
                [[2, 1], [2, 3]],
                [[2, 4], [4, 0]],
                [[4, 1], [4, 1]],
                [[4, 2], [4, 2]],
                [[4, 3], [4, 4]]
            ],
            []
        ),
        (
            ["", "Цель работы", "Основные теоретические положения", "Выполнение работы",
                 "Тестирование", "Выводы", "ПРИЛОЖЕНИЕ А", "[table]", ""],
            "LR",
            [
                [[0, 0], [0, 0]],
                [[0, 1], [0, 1]],
                [[0, 2], [0, 2]],
                [[0, 3], [0, 3]],
                [[0, 4], [0, 4]],
                [[0, 5], [0, 5]],
                [[0, 6], [2, 0]]
            ],
            []
        ),
        (
            ["", "Цель работы", "Основные теоретические положения", "Выполнение работы",
                 "Тестирование", "Выводы", "ПРИЛОЖЕНИЕ А", "[table]"],
            "LR",
            [
                [[0, 0], [0, 0]],
                [[0, 1], [0, 1]],
                [[0, 2], [0, 2]],
                [[0, 3], [0, 3]],
                [[0, 4], [0, 4]],
                [[0, 5], [0, 5]],
                [[0, 6], [1, 0]]
            ],
            []
        ),
        (
            ["", "[table]", "", ""],
            "LR",
            [
                [[0, 0], [2, 1]]
            ],
            ["Цель работы", "Основные теоретические положения", "Выполнение работы", "Тестирование", "Выводы"]
        ),
        (
            ["", "[table]", "", "Тестирование", ""],
            "LR",
            [
                [[0, 0], [2, 0]],
                [[2, 1], [2, 2]]
            ],
            ["Цель работы", "Основные теоретические положения", "Выполнение работы", "Выводы"]
        ),
        (
            ["", "Цель работы", "Основные теоретические положения", "Выполнение работы",
                 "Тестирование", "ПРИЛОЖЕНИЕ А", "[table]"],
            "LR",
            [
                [[0, 0], [0, 0]],
                [[0, 1], [0, 1]],
                [[0, 2], [0, 2]],
                [[0, 3], [0, 3]],
                [[0, 4], [0, 4]],
                [[0, 5], [1, 0]]
            ],
            ["Выводы"]
        ),
        (
            ["", "Цель работы", "Основные теоретические положения", "Выполнение работы",
                 "Тестирование", "Тестирование", "Выводы"],
            "LR",
            [
                [[0, 0], [0, 0]],
                [[0, 1], [0, 1]],
                [[0, 2], [0, 2]],
                [[0, 3], [0, 3]],
                [[0, 4], [0, 5]],
                [[0, 6], [0, 6]]
            ],
            []
        ),
        (
            ["", "Цель работы", "Основные теоретические положения", "Выполнение работы",
                 "Тестирование", "Выводы", "ПРИЛОЖЕНИЕ А", "ПРИЛОЖЕНИЕ А"],
            "LR",
            [
                [[0, 0], [0, 0]],
                [[0, 1], [0, 1]],
                [[0, 2], [0, 2]],
                [[0, 3], [0, 3]],
                [[0, 4], [0, 4]],
                [[0, 5], [0, 5]],
                [[0, 6], [0, 6]],
                [[0, 7], [0, 7]]
            ],
            []
        )
    ])
    def test_simple_file(self, tmp_path, paragraphs, file_type, expected_chapters, expected_errors):
        filename = tmp_path.as_posix() + "/test_file.docx"
        docx_document = docx.Document()
        for paragraph in paragraphs:
            if paragraph == "[table]":
                docx_document.add_table(1, 1)
            else:
                docx_document.add_paragraph(paragraph)
        docx_document.save(filename)
        docx2python_document = docx2python(filename)
        assert (expected_chapters, expected_errors) == PageCreator().makeIndices(docx2python_document, file_type)

    def test_one_section_missing(self, tmp_path):
        filename = tmp_path.as_posix() + "/test_file.docx"
        for i in range(len(self.chapters_lab)):
            docx_document = docx.Document()
            docx_document.add_paragraph("Title page")
            for j in range(len(self.chapters_lab)):
                if j != i:
                    docx_document.add_paragraph(self.chapters_lab[j])
            docx_document.save(filename)
            docx2python_document = docx2python(filename)
            (actual_chapters, actual_errors) = PageCreator().makeIndices(docx2python_document, "LR")
            assert actual_errors == [self.chapters_lab[i]]
            assert actual_chapters == [
                [[0, 0], [0, 0]],
                [[0, 1], [0, 1]],
                [[0, 2], [0, 2]],
                [[0, 3], [0, 3]],
                [[0, 4], [0, 4]]
            ]


if __name__ == '__main__':
    pytest.main()
