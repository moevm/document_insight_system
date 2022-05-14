import re

import pytest
import docx

from app.main.mse22.document.document import Document


class TestDocument:
    sections_lab = ['Цель работы', 'Основные теоретические положения',
                        'Выполнение работы', 'Тестирование', 'Выводы']
    sections_fwq = ['ЗАДАНИЕ НА ВЫПУСКНУЮ КВАЛИФИКАЦИОННУЮ РАБОТУ',
                        'КАЛЕНДАРНЫЙ ПЛАН ВЫПОЛНЕНИЯ ВЫПУСКНОЙ КВАЛИФИКАЦИОННОЙ РАБОТЫ', 'РЕФЕРАТ', 'ABSTRACT',
                        'СОДЕРЖАНИЕ', 'ОПРЕДЕЛЕНИЯ, ОБОЗНАЧЕНИЯ И СОКРАЩЕНИЯ', 'ВВЕДЕНИЕ', 'ЗАКЛЮЧЕНИЕ',
                        'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ']

    @pytest.mark.parametrize("file_type, expected_errors", [
        ("LR", sections_lab), ("FWQ", sections_fwq)
    ])
    def test_empty_file(self, tmp_path, file_type, expected_errors):
        docx_document = docx.Document()
        filename = tmp_path.as_posix() + "empty.docx"
        docx_document.save(filename)
        parsed_document = Document(docx_document, filename, file_type)
        assert parsed_document.info.author == "python-docx"
        assert parsed_document.errors == expected_errors

    def test_create_minimal_lab(self, tmp_path):
        path_string = tmp_path.as_posix()
        filename = path_string + "/minimal.docx"
        document = docx.Document()
        for section_name in self.sections_lab:
            document.add_paragraph(section_name)
        document.save(filename)
        parsed_document = Document(document, filename, "LR")
        assert parsed_document.info.author == "python-docx"
        assert parsed_document.errors == []
        assert len(parsed_document.pages) == len(self.sections_lab) + 1
        assert [pg.header for pg in parsed_document.pages] == ["Титульный лист"] + self.sections_lab
        assert parsed_document.pages[0].pageObjects == []
        for page in parsed_document.pages[1:]:
            assert len(page.pageObjects) == 1
            page_object = page.pageObjects[0]
            assert page_object.type == "paragraph"
        assert [page.pageObjects[0].text for page in parsed_document.pages[1:]] == self.sections_lab

    @pytest.mark.parametrize("filename, expected_sections", [
        ("../test_files/docx/passing-1.docx", ["Титульный лист"] + sections_lab),
        ("../test_files/docx/passing-2.docx", ["Титульный лист"] + sections_lab + ["ПРИЛОЖЕНИЕ А"]),
        ("../test_files/docx/passing-3.docx", ["Титульный лист"] + sections_lab),
        ("../test_files/docx/passing-4.docx", ["Титульный лист"] + sections_lab),
        ("../test_files/docx/passing-5.docx", ["Титульный лист"] + sections_lab + ["приложение А"])
    ])
    def test_pre_made_passing_lab_docx(self, filename, expected_sections):
        appendix_regex = re.compile("(Приложение|ПРИЛОЖЕНИЕ|приложение) [А-ЯЁа-яё]")
        docx_document = docx.Document(filename)
        parsed_document = Document(docx_document, filename, "LR")
        assert parsed_document.errors == []
        actual_sections = [page.header for page in parsed_document.pages]
        assert actual_sections == expected_sections
        for section in actual_sections:
            if section not in ["Титульный лист"] + self.sections_lab:
                print(section)
                assert appendix_regex.match(str(section))

    '''def test_pre_made_passing_lab_docx_(self):
        sections = ["Титульный лист"] + self.sections_lab
        appendix_regex = re.compile("ПРИЛОЖЕНИЕ [А-ЯЁ]")
        for i in range(1, 6):
            filename = "../test_files/docx/passing-{0}.docx".format(str(i))
            docx_document = docx.Document(filename)
            parsed_document = Document(docx_document, filename, "LR")
            assert parsed_document.errors == []
            assert len(parsed_document.pages) >= len(sections)
            assert [page.header for page in parsed_document.pages[0:len(sections)]] == sections
            for j in range(len(parsed_document.pages), len(parsed_document.pages)):
                assert appendix_regex.match(parsed_document.pages[j].header)'''

    @pytest.mark.parametrize("filename, expected_errors", [
        ("../test_files/docx/failing-1.docx", ["Выполнение работы"]),
        ("../test_files/docx/failing-2.docx", ["Цель работы"]),
        ("../test_files/docx/failing-3.docx", ["Цель работы"]),
        ("../test_files/docx/failing-4.docx", ["Основные теоретические положения"]),
        ("../test_files/docx/failing-5.docx", ['Цель работы', 'Основные теоретические положения',
                                               'Выполнение работы', 'Тестирование', 'Выводы']),
        ("../test_files/docx/failing-6.docx", ["Выводы"])
    ])
    def test_pre_made_failing_lab_docx(self, filename, expected_errors):
        parsed_document = Document(docx.Document(filename), filename, "LR")
        assert parsed_document.errors == expected_errors
        sections = ["Титульный лист"] + self.sections_lab
        for section in parsed_document.errors:
            sections.remove(section)
        actual_sections = [page.header for page in parsed_document.pages]
        for section in actual_sections:
            if re.compile("(ПРИЛОЖЕНИЕ|Приложение|приложение) [А-ЯЁа-яё]").match(section):
                actual_sections.remove(section)
        assert actual_sections == sections


if __name__ == '__main__':
    pytest.main()
