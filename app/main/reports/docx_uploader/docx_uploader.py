import re
from functools import reduce

import docx

from .core_properties import CoreProperties
from .inline_shape import InlineShape
from .paragraph import Paragraph
from .style import Style
from .table import Table, Cell
from ..pdf_document.pdf_document_manager import PdfDocumentManager


class DocxUploader:
    def __init__(self):
        self.inline_shapes = []
        self.core_properties = None
        self.headers = []
        self.chapters = []
        self.paragraphs = []
        self.tables = []
        self.file = None
        self.styled_paragraphs = None
        self.special_paragraph_indices = {}
        self.pdf_file = None
        self.count = 0
        self.first_lines = []
        self.literature_header = []
        self.headers_page = 0

    def upload(self, file, pdf_filepath=''):
        self.file = docx.Document(file)
        self.pdf_file = PdfDocumentManager(file, pdf_filepath)

    def parse(self):
        self.core_properties = CoreProperties(self.file)
        for i in range(len(self.file.inline_shapes)):
            self.inline_shapes.append(InlineShape(self.file.inline_shapes[i]))
        self.paragraphs = self.__make_paragraphs(self.file.paragraphs)
        self.parse_effective_styles()
        self.tables = self.__make_table(self.file.tables)

    def __make_paragraphs(self, paragraphs):
        tmp_paragraphs = []
        for i in range(len(paragraphs)):
            if len(paragraphs[i].text.strip()):
                tmp_paragraphs.append(Paragraph(paragraphs[i]))
        return tmp_paragraphs

    def make_chapters(self, work_type):
        if not self.chapters:
            tmp_chapters = []
            if work_type == 'VKR':
                # find headers
                header_ind = -1
                par_num = 0
                head_par_ind = -1
                for par_ind in range(len(self.styled_paragraphs)):
                    head_par_ind += 1
                    style_name = self.paragraphs[par_ind].paragraph_style_name.lower()
                    if style_name.find("heading") >= 0:
                        header_ind += 1
                        par_num = 0
                        tmp_chapters.append({"style": style_name, "text": self.styled_paragraphs[par_ind]["text"].strip(),
                                             "styled_text": self.styled_paragraphs[par_ind], "number": head_par_ind,
                                             "child": []})
                    elif header_ind >= 0:
                        par_num += 1
                        tmp_chapters[header_ind]["child"].append(
                            {"style": style_name, "text": self.styled_paragraphs[par_ind]["text"],
                             "styled_text": self.styled_paragraphs[par_ind], "number": head_par_ind})
            self.chapters = tmp_chapters
        return self.chapters

    def make_headers(self, work_type):
        if not self.headers:
            if work_type == 'VKR':
                # find first pages
                headers = [
                    {"name": "Титульный лист", "marker": False, "key": "санкт-петербургский государственный",
                     "main_character": True, "page": 0},
                    {"name": "Задание на выпускную квалификационную работу", "marker": False, "key": "задание",
                     "main_character": True, "page": 0},
                    {"name": "Календарный план", "marker": False, "key": "календарный план", "main_character": True,
                     "page": 0},
                    {"name": "Реферат", "marker": False, "key": "реферат", "main_character": False,  "page": 0},
                    {"name": "Abstract", "marker": False, "key": "abstract", "main_character": False, "page": 0},
                    {"name": "Cодержание", "marker": False, "key": "содержание", "main_character": False, "page": 0}]
                for page in range(1, self.count if self.page_counter() < 2 * len(headers) else 2 * len(headers)):
                    page_text = (self.pdf_file.get_text_on_page()[page].split("\n")[0]).lower()
                    for i in range(len(headers)):
                        if not headers[i]["marker"]:
                            if page_text.find(headers[i]["key"]) >= 0:
                                headers[i]["marker"] = True
                                headers[i]["page"] = page
                                break
                self.headers = headers
        return self.headers

    def __make_table(self, tables):
        for i in range(len(tables)):
            table = []
            for j in range(len(tables[i].rows)):
                row = []
                for k in range(len(tables[i].rows[j].cells)):
                    tmp_paragraphs = self.__make_paragraphs(tables[i].rows[j].cells[k].paragraphs)
                    row.append(Cell(tables[i].rows[j].cells[k], tmp_paragraphs))
                table.append(row)
            self.tables.append(Table(tables[i], table))
        return tables

    def find_header_page(self, work_type):
        if not self.headers_page:
            if work_type != 'VKR':
                self.headers_page = 1
                return self.headers_page
            for header in self.make_headers(work_type):
                if header["name"].find('Cодержание') >= 0:
                    if not header["page"]:
                        self.headers_page = 1
                    else:
                        self.headers_page = header["page"]
                    break
        return self.headers_page

    def find_literature_vkr(self, work_type):
        if not self.literature_header:
            for header in self.make_chapters(work_type):
                header_text = header["text"].lower()
                if header_text.find('список использованных источников') >= 0:
                    self.literature_header = header
        return self.literature_header

    def build_vkr_hierarchy(self, styles):
        indices = self.get_paragraph_indices_by_style(styles)
        tagged_indices = [{"index": 0, "level": 0}, {"index": len(self.styled_paragraphs), "level": 0}]
        for j in range(len(indices)):
            tagged_indices.extend(list(map(lambda index: {"index": index, "level": j + 1,
                                                          "text": self.styled_paragraphs[index]["text"]}, indices[j])))
        tagged_indices.sort(key=lambda dct: dct["index"])
        return tagged_indices

    def build_vkr_hierarchy_style(self, styles):
        indices = self.get_paragraph_indices_by_style(styles)
        tagged_indices = [{"index": 0, "level": 0}, {"index": len(self.styled_paragraphs), "level": 0}]
        for j in range(len(indices)):
            tagged_indices.extend(list(map(lambda index: {"index": index, "level": j + 1,
                                                          "styled_text": self.styled_paragraphs[index],
                                                          "style": self.paragraphs[index].paragraph_style_name.lower()},
                                           indices[j])))
        tagged_indices.sort(key=lambda dct: dct["index"])
        return tagged_indices

    # Parses styles once; subsequent calls have no effect, since the file itself shouldn't change
    def parse_effective_styles(self):
        if self.styled_paragraphs is not None:
            return
        self.styled_paragraphs = []
        for par in filter(lambda p: len(p.text.strip()) > 0, self.file.paragraphs):
            paragraph = {"text": par.text, "runs": []}
            for run in filter(lambda r: len(r.text.strip()) > 0, par.runs):
                paragraph["runs"].append({"text": run.text, "style": Style(run, par)})
            self.styled_paragraphs.append(paragraph)

    def unify_multiline_entities(self, first_line_regex_str):
        pattern = re.compile(first_line_regex_str)
        pars_to_delete = []
        skip_flag = False
        for i in range(len(self.styled_paragraphs) - 1):
            if skip_flag:
                skip_flag = False
                continue
            par = self.styled_paragraphs[i]
            next_par = self.styled_paragraphs[i + 1]
            if pattern.match(par["text"]):
                skip_flag = True
                par["text"] += ("\n" + next_par["text"])
                par["runs"].extend(next_par["runs"])
                pars_to_delete.append(next_par)
                continue
        for par in pars_to_delete:
            self.styled_paragraphs.remove(par)

    def get_paragraph_indices_by_style(self, style_list):
        result = []
        for template_style in style_list:
            matched_pars = []
            for i in range(len(self.styled_paragraphs)):
                par = self.styled_paragraphs[i]
                if reduce(lambda prev, run: prev and run["style"].matches(template_style), par["runs"], True):
                    matched_pars.append(i)
            result.append(matched_pars)
        return result

    def page_counter(self):
        if not self.count:
            for k, v in self.pdf_file.text_on_page.items():
                line = v[:20] if len(v) > 21 else v
                if re.search('ПРИЛОЖЕНИЕ [А-Я]', line.strip()):
                    break
                self.count += 1
                line = ''
                lines = v.split("\n")
                for i in range(len(lines)):
                    if i > 1:
                        break
                    if i > 0:
                        line += " "
                    line += lines[i].strip()
                self.first_lines.append(line.lower())
        return self.count

    def upload_from_cli(self, file):
        self.upload(file=file)

    def print_info(self):
        print(self.core_properties.to_string())
        for i in range(len(self.paragraphs)):
            print(self.paragraphs[i].to_string())

    def __str__(self):
        return self.core_properties.to_string() + '\n' + '\n'.join(
            [self.paragraphs[i].to_string() for i in range(len(self.paragraphs))])


def main(args):
    file = args.file
    uploader = DocxUploader()
    uploader.upload_from_cli(file=file)
    uploader.parse()
    uploader.print_info()
    uploader.parse_effective_styles()
