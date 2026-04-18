import re

import docx
from docx2python import docx2python

from .chapter import Chapter
from .chapter_object import ChapterObjectHeader, ChapterObjectImage, ChapterObjectTable


class ChapterCreator:
    @staticmethod
    def make_chapter_object(docx_docx2python, docx_docx, i, j, paragraph_index):
        image_pattern = r'\-{4}media/image\d+\.\D+\-{4}'
        if re.search(image_pattern, docx_docx2python.body[i][0][0][j]):
            return ChapterObjectImage('image', docx_docx.paragraphs[paragraph_index])
        else:
            return ChapterObjectHeader('paragraph', docx_docx.paragraphs[paragraph_index])

    @staticmethod
    def make_content(doc, body_index, paragraph_index):
        new_chapters = []
        index = 0
        paragraph_index += 1
        flag = False
        while (
            len(doc.body[body_index][0][0]) >= paragraph_index + 1
            and doc.body[body_index][0][0][paragraph_index] == ' '
        ):
            paragraph_index += 1
        if len(doc.body[body_index][0][0]) == paragraph_index + 1:
            body_index += 1
            paragraph_index = 0
            for elem in doc.body[body_index]:
                s = ''
                for item in elem:
                    s += item[0]
                    s += '+'
                if re.search(r'(ЗАКЛЮЧЕНИЕ|Заключение|заключение)', s):
                    flag = True
                    break
                elif (
                    re.search(r'(ВВЕДЕНИЕ|Введение|введение)', s)
                    or re.search(
                        r'(ОПРЕДЕЛЕНИЯ, ОБОЗНАЧЕНИЯ И СОКРАЩЕНИЯ|Определения, обозначения и сокращения|'
                        r'определения, обозначения и сокращения)',
                        s,
                    )
                    or len(s) == len(elem)
                ):
                    continue
                else:
                    if s.split('+')[1].strip():
                        new_chapters.append(f"{s.split('+')[0]} {s.split('+')[1]}".strip())
        else:
            for elem in doc.body[body_index][0][0]:
                if re.search(r'(ЗАКЛЮЧЕНИЕ|Заключение|заключение)', elem):
                    flag = True
                    break
            if not flag:
                return False, None, None
            for i in range(paragraph_index, len(doc.body[body_index][0][0])):
                if (
                    re.search(r'(ВВЕДЕНИЕ|Введение|введение)', doc.body[body_index][0][0][i])
                    or re.search(
                        r'(ОПРЕДЕЛЕНИЯ, ОБОЗНАЧЕНИЯ И СОКРАЩЕНИЯ|Определения, обозначения и сокращения|'
                        r'определения, обозначения и сокращения)',
                        doc.body[body_index][0][0][i],
                    )
                    or doc.body[body_index][0][0][i] == ''
                ):
                    continue
                elif re.search(r'(ЗАКЛЮЧЕНИЕ|Заключение|заключение)', doc.body[body_index][0][0][i]):
                    index = i + 1
                    break
                else:
                    new_chapters.append(
                        doc.body[body_index][0][0][i][
                            re.search(r'[\w. ]{2,}', doc.body[body_index][0][0][i]).start() : re.search(
                                r'[\w. ]{2,}', doc.body[body_index][0][0][i]
                            ).end()
                        ]
                    )
        if flag:
            return True, new_chapters, index
        else:
            return False, None, None

    def create_chapter_objects(self, path, file_type):
        docx_docx2python = docx2python(path)
        indices, errors = self.make_indices(docx_docx2python, file_type)
        docx_docx = docx.Document(path)

        table_index = 0
        paragraph_index = 0
        page_objects = []
        current_page_object = []
        for i in range(0, len(indices)):
            chapter = docx_docx.paragraphs[paragraph_index].text if i != 0 else 'Титульный лист'
            for j in range(indices[i][0][0], indices[i][1][0] + 1):
                if len(docx_docx2python.body[j]) > 1:
                    current_page_object.append(ChapterObjectTable('table', docx_docx.tables[table_index]))
                    table_index += 1

                else:
                    if indices[i][0][0] == indices[i][1][0]:
                        for _ in range(indices[i][0][1], indices[i][1][1] + 1):
                            current_page_object.append(
                                self.make_chapter_object(docx_docx2python, docx_docx, j, _, paragraph_index)
                            )
                            paragraph_index += 1

                    elif j == indices[i][0][0] and j != indices[i][1][0]:
                        for _ in range(indices[i][0][1], len(docx_docx2python.body[j][0][0])):
                            current_page_object.append(
                                self.make_chapter_object(docx_docx2python, docx_docx, j, _, paragraph_index)
                            )
                            paragraph_index += 1

                    elif j != indices[i][0][0] and j != indices[i][1][0]:
                        for _ in range(0, len(docx_docx2python.body[j][0][0])):
                            current_page_object.append(
                                self.make_chapter_object(docx_docx2python, docx_docx, j, _, paragraph_index)
                            )
                            paragraph_index += 1

                    elif j == indices[i][1][0]:
                        for _ in range(0, indices[i][1][1] + 1):
                            current_page_object.append(
                                self.make_chapter_object(docx_docx2python, docx_docx, j, _, paragraph_index)
                            )
                            paragraph_index += 1

            page_objects.append(Chapter(chapter, current_page_object))
            current_page_object = []

        return page_objects, errors

    def make_indices(self, parsed_doc, file_type, chapter_names=None):
        # 'LR':
        #     chapters = ['Цель работы', 'Основные теоретические положения',
        #                 'Выполнение работы', 'Тестирование', 'Выводы']
        # else:
        #     chapters = ['ЗАДАНИЕ НА ВЫПУСКНУЮ КВАЛИФИКАЦИОННУЮ РАБОТУ',
        #                 'КАЛЕНДАРНЫЙ ПЛАН ВЫПОЛНЕНИЯ ВЫПУСКНОЙ КВАЛИФИКАЦИОННОЙ РАБОТЫ', 'РЕФЕРАТ', 'ABSTRACT',
        #                 'СОДЕРЖАНИЕ', 'ОПРЕДЕЛЕНИЯ, ОБОЗНАЧЕНИЯ И СОКРАЩЕНИЯ', 'ВВЕДЕНИЕ', 'ЗАКЛЮЧЕНИЕ',
        #                 'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ']
        i_start = [0, 0]
        doc_chapters = []  #
        errors = []
        cur_index = 0
        cur_chapter = 0
        application_pattern = r'(Приложение|ПРИЛОЖЕНИЕ|приложение) [А-ЯЁа-яё]'
        application_find_index = 0
        find_index = 0
        while cur_chapter != len(chapter_names):
            if not parsed_doc.body:
                # пустой документ
                errors = chapter_names.copy()
                break

            if cur_index == len(parsed_doc.body):
                # ???
                if chapter_names[cur_chapter] != 'определения, обозначения и сокращения':
                    errors.append(chapter_names[cur_chapter])
                cur_chapter += 1
                cur_index = doc_chapters[-1][0][0] if doc_chapters else 0
                if cur_chapter == len(chapter_names):
                    break

            if len(parsed_doc.body[cur_index]) > 1:  # table
                cur_index += 1
            elif (
                chapter_names[cur_chapter] in parsed_doc.body[cur_index][0][0][find_index:]
                or chapter_names[cur_chapter].lower() in parsed_doc.body[cur_index][0][0][find_index:]
                and file_type != 'LR'
                or chapter_names[cur_chapter].capitalize() in parsed_doc.body[cur_index][0][0][find_index:]
            ):
                if chapter_names[cur_chapter] in parsed_doc.body[cur_index][0][0][find_index:]:
                    find_chapter = chapter_names[cur_chapter]
                elif chapter_names[cur_chapter].lower() in parsed_doc.body[cur_index][0][0][find_index:]:
                    find_chapter = chapter_names[cur_chapter].lower()
                else:
                    find_chapter = chapter_names[cur_chapter].capitalize()
                if chapter_names[cur_chapter] == 'СОДЕРЖАНИЕ':
                    is_correct, new_chapters, index = self.make_content(
                        parsed_doc, cur_index, parsed_doc.body[cur_index][0][0].index(find_chapter)
                    )
                    if is_correct:
                        chapter_names = (
                            chapter_names[: cur_chapter + 3] + new_chapters + chapter_names[cur_chapter + 3 :]
                        )

                        if not doc_chapters:
                            doc_chapters.append(
                                [i_start, [cur_index, parsed_doc.body[cur_index][0][0].index(find_chapter) - 1]]
                            )
                        else:
                            doc_chapters.append(
                                [
                                    [doc_chapters[-1][1][0], doc_chapters[-1][1][1] + 1],
                                    [cur_index, parsed_doc.body[cur_index][0][0].index(find_chapter) - 1],
                                ]
                            )

                        find_index = index
                        i_start = [doc_chapters[-1][1][0], doc_chapters[-1][1][1] + 1]
                        cur_chapter += 1
                        continue

                tmp_i_end = [0, 0]
                tmp_i_start = i_start.copy()
                if parsed_doc.body[cur_index][0][0].index(find_chapter, find_index) == 0 and cur_index > 0:
                    tmp_i_end[0], tmp_i_end[1] = cur_index - 1, len(parsed_doc.body[cur_index - 1][0][0]) - 1
                    i_start[0], i_start[1] = tmp_i_end[0] + 1, 0
                else:
                    tmp_i_end[0], tmp_i_end[1] = (
                        cur_index,
                        parsed_doc.body[cur_index][0][0].index(find_chapter, find_index) - 1,
                    )
                    i_start[0], i_start[1] = tmp_i_end[0], tmp_i_end[1] + 1
                doc_chapters.append([tmp_i_start, tmp_i_end])
                cur_chapter += 1
            else:
                cur_index += 1

        for i in range(cur_index, len(parsed_doc.body)):
            if len(parsed_doc.body[i][0][0]) > 1:
                for j in range(application_find_index, len(parsed_doc.body[i][0][0])):
                    if re.search(application_pattern, parsed_doc.body[i][0][0][j]):
                        if not doc_chapters:
                            doc_chapters.append([[0, 0], [i, j - 1]])
                        else:
                            doc_chapters.append([[doc_chapters[-1][1][0], doc_chapters[-1][1][1] + 1], [i, j - 1]])
                        application_find_index = j - 1
            cur_index += 1
        if not doc_chapters:
            if parsed_doc.body:  # empty file
                doc_chapters.append([[0, 0], [len(parsed_doc.body) - 1, len(parsed_doc.body[-1][0][0]) - 1]])
        else:
            doc_chapters.append(
                [
                    [doc_chapters[-1][1][0], doc_chapters[-1][1][1] + 1],
                    [len(parsed_doc.body) - 1, len(parsed_doc.body[-1][0][0]) - 1],
                ]
            )
        return doc_chapters, errors
