import fitz
from docx import Document
from ..base_check import BaseReportCriterion, answer

default_font_size: int = 12 # Значение размера шрифта по умолчанию
default_line_spacing: float = 1.0 # Значение межстрочного интервала по умолчанию
chars_per_line: int = 20 # Примерное кол-во символов в строке ячейки
cell_padding = 5 # Примерный размер отступов от границ ячейки
row_padding = 2 # Примерное расстояние между строками таблиц
#С данными значениями погрешность 4-5 процентов


class ReportTablePercentageCheck(BaseReportCriterion):
    label = 'Проверка процентного соотношения таблиц в документе'
    _description = 'Проверяет, что таблицы занимают не более установленного процента от объёма документа'
    id = 'report_table_percentage_check'


    def __init__(self, file_info, has_application=True, max_percentage=15, max_pages_table = 2):
        super().__init__(file_info)
        self._max_percentage = max_percentage
        self._hasApplication = has_application
        self._large_tables = []
        self._max_pages_table = max_pages_table

    def get_font_size(self, run, paragraph) -> float:
        """Функция получения размера шрифта"""
        if run.font.size:
            return run.font.size.pt
        elif run.style and run.style.font.size:
            return run.style.font.size.pt
        elif paragraph.style and paragraph.style.font.size:
            return paragraph.style.font.size.pt
        else:
            return default_font_size

    def height_cell(self, cell) -> float:
        """Функция получения высоты ячейки"""
        total_height = 0

        for paragraph in cell.paragraphs:
            line_count = 1

            text_length = len(paragraph.text)
            if text_length > 0:
                line_count = max(1, (text_length + chars_per_line - 1) // chars_per_line)

            max_font_size = default_font_size
            for run in paragraph.runs:
                font_size = self.get_font_size(run, paragraph)
                max_font_size = max(max_font_size, font_size)

            line_spacing = default_line_spacing

            paragraph_height = line_count * max_font_size * line_spacing
            total_height += paragraph_height

            if paragraph.paragraph_format.space_after:
                total_height += paragraph.paragraph_format.space_after.pt
            if paragraph.paragraph_format.space_before:
                total_height += paragraph.paragraph_format.space_before.pt

        total_height += 2 * cell_padding

        return total_height

    def height_table(self, table) -> float:
        """Функция получения высоты таблицы"""
        total_height = 0
        for row in table.rows:
            if row.height:
                total_height += row.height.pt
            else:
                heights = []
                for cell in row.cells:
                    heights.append(self.height_cell(cell))
                total_height += max(heights) if heights else 0

            total_height += row_padding

        return total_height

    def get_percent_of_tables(self) -> float:
        """Функция получение процента таблиц в документе"""
        doc_docx = self.file.file

        page_count = self.file.page_counter()
        if page_count == 0:
            return 0

        section = doc_docx.sections[0]
        page_height = section.page_height.pt
        total_height = page_height * page_count

        if total_height == 0:
            return 0

        tables_height = 0

        end_index = len(doc_docx.tables)

        if self._hasApplication:
            end_index = self.find_table_index_after_text('ПРИЛОЖЕНИЕ')
            end_index = end_index if end_index is not None else len(doc_docx.tables)

        for table_index, table in enumerate(doc_docx.tables):
            if table_index < end_index:
                table_height = self.height_table(table)
                if table_height > page_height * self._max_pages_table:
                    self._large_tables.append(table_index)
                tables_height += table_height

        percentage = (tables_height / total_height) * 100
        return percentage


    def find_table_index_after_text(self, target_text: str) -> int | None:
        """Функция находит первый индекс таблицы после target_text, если не найден, то вернет None"""
        doc_docx = self.file.file
        for i, paragraph in enumerate(doc_docx.paragraphs):
            if target_text in paragraph.text:

                para_elem = paragraph._element
                following_tables = para_elem.xpath('./following-sibling::w:tbl')

                for table_elem in following_tables:

                    for index, table in enumerate(doc_docx.tables):
                        if table._element == table_elem:
                            return index

        return None

    def find_tables_indexes_between_text(self,start_text: str, end_text: str | None = None):
        """Функция нахождения индексов начала и конца таблиц между двумя текстами"""
        start_index = self.find_table_index_after_text(start_text)
        end_index = self.find_table_index_after_text(end_text)

        return start_index, end_index

    def check(self):
        try:
            percent_tables = self.get_percent_of_tables()

            message = ''

            if percent_tables <= self._max_percentage and len(self._large_tables) == 0:
                return answer(
                    True,
                    "Пройдена!"
                )

            elif len(self._large_tables) != 0:
                message = (f'Есть таблицы, занимающие более {self._max_pages_table} страниц:\n'
                           f'<ul>\n'
                           + ''.join(f'<li>Таблица {index_table}</li>\n' for index_table in self._large_tables)
                           + '</ul>')
            else:
                message += (
                    f'Таблицы занимают {percent_tables:.1f}% документа, '
                    f'что превышает допустимые {self._max_percentage}%. '
                    'Рекомендации: \n'
                    '<ul>\n'
                    '    <li>Уменьшите количество таблиц в основном тексте документа;</li>\n'
                    '    <li>Перенесите вспомогательные таблицы в приложения;</li>\n'
                    '    <li>Сократите объем данных в таблицах, оставив только самую важную информацию.</li>\n'
                    '</ul>'
                )

            return answer(
                False,
                message
            )

        except Exception as e:
            return answer(False, f"Ошибка при проверке процентного соотношения таблиц: {str(e)}")

